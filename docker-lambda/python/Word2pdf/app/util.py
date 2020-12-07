import time
import datetime
import os
import glob
import subprocess
import boto3
from copy import deepcopy
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

# clearDirectory
def clearDirectory(directory):
  files = glob.glob(f'{directory}*')
  for f in files:
    os.remove(f)

# docx_replace
def docx_replace(wdoc, regex, replace):

  for p in wdoc.paragraphs:
    if regex.search(p.text):
      inline = p.runs
      for i in range(len(inline)):
        if regex.search(inline[i].text):
          text = regex.sub(replace, inline[i].text)
          inline[i].text = text

  for table in wdoc.tables:
    for row in table.rows:
      for cell in row.cells:
        docx_replace(cell, regex, replace)

# docx_fill_data
def docx_fill_data(wdoc, data):
  
  items = data['items']
  total = data['total']

  table = wdoc.tables[1]

  def set_row_height(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), "250")
    trHeight.set(qn('w:hRule'), "atLeast")
    trPr.append(trHeight)

  int_format = '{:,}'
  decimal_format = '{:,.2f}'
  for item in items:
    row = table.add_row()
    set_row_height(row)
    cells = row.cells
    cells[0].text = int_format.format(item['quantity'])
    cells[1].text = str(item['description'])
    cells[2].text = decimal_format.format(item['unitprice'])
    cells[3].text = decimal_format.format(item['linetotal'])

    for x in range(4):
      if (x != 1):
        cells[x].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

      cells[x].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
      font = cells[x].paragraphs[0].runs[0].font
      font.size = Pt(8)

  row = table.add_row()
  cells = row.cells
  cells[2].text = str('Total')
  cells[3].text = decimal_format.format(total)

  for x in range(2, 4):
    cell = cells[x]
    font = cell.paragraphs[0].runs[0].font
    font.bold = True
    font.size = Pt(8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

  background = parse_xml(r'<w:shd {} w:fill="E7EEEE"/>'.format(nsdecls('w')))
  index = 0
  for row in table.rows:
    index += 1
    if (index == 1 or index % 2 == 0):
      continue

    for cell in row.cells:
      cell._tc.get_or_add_tcPr().append(deepcopy(background))

# doc2pdf
def doc2pdf(word_file, pdf_file):  
  cmd = [f'/app/word-pdf/word-pdf']
  print(cmd)

  p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
  p.wait()
  stdout, stderr = p.communicate()

# upload2s3
def upload2s3(bucket, result_pdf_file):
  ts = datetime.datetime.fromtimestamp(time.time()).strftime(r'%Y-%m-%d-%H-%M-%S')
  object_name = f'G-{ts}.pdf'

  s3 = boto3.resource('s3')
  s3.meta.client.upload_file(result_pdf_file, bucket, object_name)
